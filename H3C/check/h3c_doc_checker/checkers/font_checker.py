"""字体格式检查模块"""
import re
from typing import Dict, List, Any, Optional, Tuple
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt
from h3c_doc_checker.utils import CheckResult, get_paragraph_style_name

class FontChecker:
    """字体格式检查器类"""
    
    def __init__(self, doc: Document, rules: Dict[str, Any]):
        """
        初始化字体检查器
        
        Args:
            doc: Word文档对象
            rules: 字体检查规则，格式可以是：
                  {
                      "heading_font_rules": {
                          "Heading 4": {
                              "chinese_font": "黑体",
                              "english_font": "Arial",
                              "font_size": 10.5  # 五号字对应10.5pt
                          }
                      },
                      "content_font_rules": {
                          "chinese_fonts": ["宋体", "黑体", "楷体", "仿宋"],  # 支持多种中文字体
                          "english_fonts": ["Arial", "Times New Roman"],  # 支持多种英文字体
                          "font_size": 10.5  # 五号字对应10.5pt
                      },
                      "mixed_font_patterns": [
                          "Ubuntu", "Linux", "Windows", "Docker", "ResNet", "INT8", "QPS", 
                          "x86", "ARM", "CPU", "GPU", "NPU", "TPU", "API", "JSON", "HTTPS", 
                          "HTTP", "FTP", "SDK", "AI", "ML", "DL", "TensorFlow", "PyTorch",
                          "ResNet50", "xxx", "results", "results.json", "#", "&", "-"
                      ]
                  }
                  或者是包含font_rules字段的完整配置：
                  {
                      "font_rules": {
                          "heading_font_rules": {...},
                          "content_font_rules": {...},
                          "mixed_font_patterns": [...]
                      }
                  }
        """
        self.doc = doc
        
        # 支持两种格式的配置:
        # 1. 直接传入font_rules
        # 2. 传入包含font_rules字段的配置
        if "font_rules" in rules:
            self.rules = rules["font_rules"]
        else:
            self.rules = rules
        
        # 处理向后兼容性：将单一字体配置转换为列表格式
        content_rules = self.rules.get("content_font_rules", {})
        if "chinese_font" in content_rules and "chinese_fonts" not in content_rules:
            content_rules["chinese_fonts"] = [content_rules["chinese_font"]]
        if "english_font" in content_rules and "english_fonts" not in content_rules:
            content_rules["english_fonts"] = [content_rules["english_font"]]
            
        # 获取混合字体模式列表
        self.mixed_font_patterns = self.rules.get("mixed_font_patterns", [])
        
        # 如果没有配置混合字体模式，添加默认的技术词汇
        if not self.mixed_font_patterns:
            self.mixed_font_patterns = [
                "Ubuntu", "Linux", "Windows", "Docker", "ResNet", "INT8", "QPS", 
                "x86", "ARM", "CPU", "GPU", "NPU", "TPU", "API", "JSON", "HTTPS", 
                "HTTP", "FTP", "SDK", "AI", "ML", "DL", "TensorFlow", "PyTorch",
                "ResNet50", "xxx", "results", "results.json", "#", "&", "-"
            ]
        
    def _convert_size_to_pt(self, size_name: str) -> float:
        """将中文字号名称转换为磅值"""
        size_mapping = {
            "初号": 42,
            "小初": 36,
            "一号": 26,
            "小一": 24,
            "二号": 22,
            "小二": 18,
            "三号": 16,
            "小三": 15,
            "四号": 14,
            "小四": 12,
            "五号": 10.5,
            "小五": 9,
            "六号": 7.5,
            "小六": 6.5,
            "七号": 5.5,
            "八号": 5
        }
        return size_mapping.get(size_name, 10.5)  # 默认五号
        
    def _is_chinese_char(self, char: str) -> bool:
        """判断单个字符是否为中文"""
        return bool(re.match(r'[\u4e00-\u9fff]', char))
        
    def _is_english_or_number_char(self, char: str) -> bool:
        """判断单个字符是否为英文字母或数字"""
        return bool(re.match(r'[a-zA-Z0-9]', char))
        
    def _get_expected_font(self, char: str, style_rules: Dict[str, Any]) -> Optional[str]:
        """
        根据字符类型和样式规则获取期望的字体
        
        Args:
            char: 字符
            style_rules: 样式规则
            
        Returns:
            期望的字体名称
        """
        if self._is_chinese_char(char):
            return style_rules.get("chinese_font")
        elif self._is_english_or_number_char(char):
            return style_rules.get("english_font")
        return None
        
    def _get_font_from_run(self, run: Run, is_chinese: bool = False) -> Optional[str]:
        """
        从run中获取字体名称
        
        Args:
            run: Run对象
            is_chinese: 是否为中文字符
            
        Returns:
            字体名称
        """
        # 首先检查XML属性（更精确的方式）
        run_props = run._element.rPr
        if run_props is not None:
            for child in run_props:
                tag = child.tag.split('}')[-1]
                if tag == 'rFonts':
                    # 根据字符类型选择不同的字体属性
                    if is_chinese:
                        east_asia = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                        if east_asia:
                            return east_asia
                    else:
                        ascii_font = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                        if ascii_font:
                            return ascii_font
        
        # 如果XML中没有找到相应设置，检查run.font直接设置
        # 注意：对于混合中英文的run，这个设置可能不准确
        if run.font.name:
            # 如果是中文字符但没有找到eastAsia设置，尝试查找段落样式中的eastAsia设置
            if is_chinese:
                para = run._parent
                if para and para.style:
                    # 尝试从段落样式的XML中查找eastAsia设置
                    if para._element.pPr is not None:
                        for ppr_child in para._element.pPr:
                            if ppr_child.tag.endswith('rPr'):
                                for rpr_child in ppr_child:
                                    if rpr_child.tag.endswith('rFonts'):
                                        east_asia = rpr_child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                                        if east_asia:
                                            return east_asia
            
            return run.font.name
            
        # 检查段落样式
        para = run._parent
        if para and para.style and para.style.font:
            # 如果是中文字符，尝试获取eastAsia字体
            if is_chinese and hasattr(para.style.font, 'name_fareast') and para.style.font.name_fareast:
                return para.style.font.name_fareast
            return para.style.font.name
            
        return None
        
    def _get_effective_font_size(self, para: Paragraph, run: Optional[Run] = None) -> Optional[float]:
        """
        获取有效的字体大小，考虑继承关系
        优先级：Run直接设置 > 段落样式 > 继承样式
        """
        font_size = None
        
        # 首先检查run的直接设置
        if run is not None:
            # 检查run的XML属性
            run_props = run._element.rPr
            if run_props is not None:
                for child in run_props:
                    tag = child.tag.split('}')[-1]
                    if tag == 'sz':
                        # sz属性值是字号的两倍
                        sz_val = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        if sz_val:
                            font_size = float(sz_val) / 2
                            break
            
            # 如果XML中没有找到，检查font对象
            if font_size is None and run.font.size:
                font_size = run.font.size.pt
        
        # 如果run没有设置字体大小，检查段落样式
        if font_size is None and para.style:
            # 检查段落样式的XML属性
            if para._element.pPr is not None:
                for ppr_child in para._element.pPr:
                    if ppr_child.tag.endswith('rPr'):
                        for rpr_child in ppr_child:
                            if rpr_child.tag.endswith('sz'):
                                sz_val = rpr_child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if sz_val:
                                    font_size = float(sz_val) / 2
                                    break
            
            # 如果XML中没有找到，检查style对象
            if font_size is None and para.style.font and para.style.font.size:
                font_size = para.style.font.size.pt
        
        return font_size
        
    def check_heading_fonts(self) -> List[CheckResult]:
        """检查标题字体，分别处理中文和英文字符"""
        heading_rules = self.rules.get("heading_font_rules", {})
        
        if not heading_rules:
            return [CheckResult(
                type="标题字体检查",
                passed=True,
                message="没有标题字体检查规则",
                details={"location": "配置文件"}
            )]
            
        results = []
            
        for para_idx, para in enumerate(self.doc.paragraphs, 1):
            style_name = get_paragraph_style_name(para)
            
            if style_name in heading_rules:
                rule = heading_rules[style_name]
                para_text = para.text.strip()
                if not para_text:
                    continue
                    
                # 检查字号
                expected_size = rule.get("font_size")
                if isinstance(expected_size, str):
                    expected_size = self._convert_size_to_pt(expected_size)
                    
                actual_size = self._get_effective_font_size(para)
                    
                size_error = None
                if expected_size:
                    if actual_size:
                        if abs(actual_size - expected_size) > 0.1:  # 允许0.1pt的误差
                            size_error = f"期望字号: {expected_size}pt, 实际字号: {actual_size}pt"
                    else:
                        size_error = f"期望字号: {expected_size}pt, 实际字号: 未设置"
                
                # 按字符类型分别检查字体
                chinese_errors = []
                english_errors = []
                
                for run in para.runs:
                    run_text = run.text
                    if not run_text.strip():
                        continue
                        
                    # 分析每个字符
                    for char in run_text:
                        if not char.strip():  # 跳过空白字符
                            continue
                            
                        expected_font = self._get_expected_font(char, rule)
                        if not expected_font:
                            continue
                            
                        is_chinese = self._is_chinese_char(char)
                        actual_font = self._get_font_from_run(run, is_chinese)
                        
                        # 比较字体
                        if actual_font and actual_font != expected_font:
                            if is_chinese:
                                chinese_errors.append(f"字符 '{char}' 期望字体: {expected_font}, 实际字体: {actual_font}")
                            else:
                                english_errors.append(f"字符 '{char}' 期望字体: {expected_font}, 实际字体: {actual_font}")
                
                # 生成检查结果
                if chinese_errors or english_errors or size_error:
                    # 精简标题字体检查输出
                    error_msg = f"标题 '{para_text}' (样式: {style_name}) 字体格式错误:\n"
                    
                    if chinese_errors:
                        error_msg += f"中文字体问题: 共 {len(set(chinese_errors))} 处\n"
                        # 最多显示1个示例
                        if chinese_errors:
                            error_msg += f"  示例: {list(set(chinese_errors))[0]}\n"
                            
                    if english_errors:
                        error_msg += f"英文/数字字体问题: 共 {len(set(english_errors))} 处\n"
                        # 最多显示1个示例
                        if english_errors:
                            error_msg += f"  示例: {list(set(english_errors))[0]}\n"
                            
                    if size_error:
                        error_msg += f"字号问题: {size_error}"
                    
                    results.append(CheckResult(
                        type="标题字体检查",
                        passed=False,
                        message=error_msg,
                        details={
                            "location": f"标题: {para_text}",
                            "style": style_name,
                            "chinese_errors_count": len(set(chinese_errors)),
                            "english_errors_count": len(set(english_errors)),
                            "size_error": size_error
                        }
                    ))
                else:
                    results.append(CheckResult(
                        type="标题字体检查",
                        passed=True,
                        message=f"标题 '{para_text}' (样式: {style_name}) 字体格式正确",
                        details={"location": f"标题: {para_text}", "style": style_name}
                    ))
        
        if not results:
            results.append(CheckResult(
                type="标题字体检查",
                passed=True,
                message="未找到需要检查的标题",
                details={"location": "整个文档"}
            ))
            
        return results
        
    def check_content_fonts(self) -> List[CheckResult]:
        """检查正文字体，分别处理中文和英文字符"""
        results = []
        content_rules = self.rules.get("content_font_rules", {})
        
        if not content_rules:
            return [CheckResult(
                type="正文字体检查",
                passed=True,
                message="没有正文字体检查规则",
                details={"location": "配置文件"}
            )]
            
        # 支持多种可接受的字体
        chinese_fonts = content_rules.get("chinese_fonts", [])
        english_fonts = content_rules.get("english_fonts", [])
        
        # 向后兼容：如果没有设置多字体，则尝试使用单一字体配置
        if not chinese_fonts and "chinese_font" in content_rules:
            chinese_fonts = [content_rules["chinese_font"]]
        if not english_fonts and "english_font" in content_rules:
            english_fonts = [content_rules["english_font"]]
            
        expected_size = content_rules.get("font_size")
        
        # 转换字号名称为磅值
        if isinstance(expected_size, str):
            expected_size = self._convert_size_to_pt(expected_size)
        
        if not chinese_fonts and not english_fonts and not expected_size:
            return [CheckResult(
                type="正文字体检查",
                passed=True,
                message="未设置正文字体规则",
                details={"location": "配置文件"}
            )]
            
        # 获取expected_titles列表
        title_rules = self.rules.get("title_rules", {})
        expected_titles = title_rules.get("expected_titles", [])
        
        if not expected_titles:
            return [CheckResult(
                type="正文字体检查",
                passed=True,
                message="没有定义expected_titles，跳过正文字体检查",
                details={"location": "配置文件"}
            )]
            
        chinese_errors = []
        english_errors = []
        size_errors = []
        
        # 获取所有expected_titles的文本
        expected_title_texts = [title.get("text", "").strip() for title in expected_titles]
        
        # 将混合字体模式关键词转换为小写，用于不区分大小写的匹配
        mixed_font_patterns_lower = [pattern.lower() for pattern in self.mixed_font_patterns]
        
        # 遍历文档段落，找到expected_titles下的正文段落
        for para_idx, para in enumerate(self.doc.paragraphs, 1):
            para_text = para.text.strip()
            style_name = get_paragraph_style_name(para)
            
            # 跳过标题段落
            if style_name and style_name.startswith(("Heading", "标题")):
                continue
                
            if not para_text:
                continue
                
            # 检查当前段落是否在某个expected_title下面
            is_under_expected_title = self._is_paragraph_under_expected_title(
                para_idx, expected_title_texts
            )
            
            if not is_under_expected_title:
                continue  # 跳过不在expected_titles下的正文
                
            # 检查是否是允许混合字体的段落（大小写不敏感）
            is_mixed_font_paragraph = False
            para_text_lower = para_text.lower()
            
            # 首先检查特殊字符
            special_chars = ['#', '$', '&', '-']
            for char in special_chars:
                if char in para_text:
                    is_mixed_font_paragraph = True
                    break
                    
            # 如果没有匹配到特殊字符，再检查关键词
            if not is_mixed_font_paragraph:
                for pattern_lower in mixed_font_patterns_lower:
                    if pattern_lower in para_text_lower:
                        is_mixed_font_paragraph = True
                        break
                    
            # 检查字号（段落级别检查，避免重复）
            if expected_size:
                actual_size = self._get_effective_font_size(para)
                if actual_size:
                    if abs(actual_size - expected_size) > 0.1:  # 允许0.1pt的误差
                        size_errors.append({
                            "paragraph": para_idx,
                            "text": para_text[:30] + "..." if len(para_text) > 30 else para_text,
                            "expected": expected_size,
                            "actual": actual_size
                        })
                else:
                    size_errors.append({
                        "paragraph": para_idx,
                        "text": para_text[:30] + "..." if len(para_text) > 30 else para_text,
                        "expected": expected_size,
                        "actual": None
                    })
                
            for run in para.runs:
                run_text = run.text
                if not run_text.strip():
                    continue
                    
                # 检查run中是否包含特殊字符，如果有，则将整个run视为混合字体
                run_has_special_char = False
                for char in special_chars:
                    if char in run_text:
                        run_has_special_char = True
                        break
                
                # 分析每个字符
                for char in run_text:
                    if not char.strip():  # 跳过空白字符
                        continue
                        
                    is_chinese = self._is_chinese_char(char)
                    is_english = self._is_english_or_number_char(char)
                    
                    # 如果是混合字体段落或run包含特殊字符，对中文字符在英文字体中的情况进行特殊处理
                    if (is_mixed_font_paragraph or run_has_special_char) and is_chinese:
                        actual_font = self._get_font_from_run(run, is_chinese)
                        # 如果是英文字体，但在混合字体段落中，允许这种情况
                        if actual_font and actual_font in english_fonts:
                            continue
                    
                    if is_chinese and chinese_fonts:
                        actual_font = self._get_font_from_run(run, True)
                        # 检查字体是否在允许的范围内
                        if actual_font and actual_font not in chinese_fonts:
                            chinese_errors.append({
                                "paragraph": para_idx,
                                "char": char,
                                "text": run_text[:20] + "..." if len(run_text) > 20 else run_text,
                                "expected": "、".join(chinese_fonts),  # 显示所有允许的字体
                                "actual": actual_font,
                                "is_mixed_font_paragraph": is_mixed_font_paragraph
                            })
                    elif is_english and english_fonts:
                        actual_font = self._get_font_from_run(run, False)
                        # 检查字体是否在允许的范围内
                        if actual_font and actual_font not in english_fonts:
                            english_errors.append({
                                "paragraph": para_idx,
                                "char": char,
                                "text": run_text[:20] + "..." if len(run_text) > 20 else run_text,
                                "expected": "、".join(english_fonts),  # 显示所有允许的字体
                                "actual": actual_font
                            })
        
        # 生成检查结果
        if chinese_errors or english_errors or size_errors:
            # 统计各类错误数量而不是详细列出
            error_counts = {
                "chinese": len(chinese_errors),
                "english": len(english_errors),
                "size": len(size_errors)
            }
            
            # 创建详细的错误消息
            error_msg = "正文字体格式错误:\n\n"
            
            # 按段落组织错误信息
            paragraph_errors = {}
            
            # 收集每个段落的所有错误
            for error in chinese_errors:
                para_idx = error['paragraph']
                if para_idx not in paragraph_errors:
                    paragraph_errors[para_idx] = {
                        'chinese': [], 
                        'english': [], 
                        'size': [], 
                        'text': error['text'],
                        'is_mixed_font_paragraph': error.get('is_mixed_font_paragraph', False)
                    }
                paragraph_errors[para_idx]['chinese'].append(error)
                if error.get('is_mixed_font_paragraph', False):
                    paragraph_errors[para_idx]['is_mixed_font_paragraph'] = True
                
            for error in english_errors:
                para_idx = error['paragraph']
                if para_idx not in paragraph_errors:
                    paragraph_errors[para_idx] = {
                        'chinese': [], 
                        'english': [], 
                        'size': [], 
                        'text': error['text'],
                        'is_mixed_font_paragraph': False
                    }
                paragraph_errors[para_idx]['english'].append(error)
                
            for error in size_errors:
                para_idx = error['paragraph']
                if para_idx not in paragraph_errors:
                    paragraph_errors[para_idx] = {
                        'chinese': [], 
                        'english': [], 
                        'size': [], 
                        'text': error['text'],
                        'is_mixed_font_paragraph': False
                    }
                paragraph_errors[para_idx]['size'].append(error)
            
            # 按段落显示错误
            for para_idx, errors in sorted(paragraph_errors.items()):
                # 找到段落所在的标题
                title = self._find_parent_title(para_idx)
                title_text = f"[{title}]" if title else "[未知标题]"
                
                # 添加混合字体段落的标记
                mixed_font_note = " (混合字体段落)" if errors.get('is_mixed_font_paragraph', False) else ""
                
                error_msg += f"第{para_idx}段 {title_text}{mixed_font_note}\n"
                error_msg += f"段落内容: {errors['text']}\n"
                
                if errors['chinese']:
                    error_msg += f"  中文字体错误: {len(errors['chinese'])} 处\n"
                    for error in errors['chinese'][:3]:  # 只显示前3个错误
                        error_msg += f"    字符 '{error['char']}': 期望字体: {error['expected']}, 实际字体: {error['actual']}\n"
                    if len(errors['chinese']) > 3:
                        error_msg += f"    ... 等更多错误 ...\n"
                
                if errors['english']:
                    error_msg += f"  英文字体错误: {len(errors['english'])} 处\n"
                    for error in errors['english'][:3]:
                        error_msg += f"    字符 '{error['char']}': 期望字体: {error['expected']}, 实际字体: {error['actual']}\n"
                    if len(errors['english']) > 3:
                        error_msg += f"    ... 等更多错误 ...\n"
                
                if errors['size']:
                    error_msg += f"  字号错误:\n"
                    for error in errors['size']:
                        if error['actual']:
                            error_msg += f"    期望字号: {error['expected']}pt, 实际字号: {error['actual']}pt\n"
                        else:
                            error_msg += f"    期望字号: {error['expected']}pt, 实际字号: 未设置\n"
                
                error_msg += "\n"  # 段落之间添加空行
                    
            results.append(CheckResult(
                type="正文字体检查",
                passed=False,
                message=error_msg,
                details={
                    "location": "expected_titles下的正文段落",
                    "error_counts": error_counts,
                    # 保存完整错误信息但不直接显示
                    "error_samples": {
                        "chinese": chinese_errors[:3] if chinese_errors else [],
                        "english": english_errors[:3] if english_errors else [],
                        "size": size_errors[:3] if size_errors else []
                    }
                }
            ))
        else:
            results.append(CheckResult(
                type="正文字体检查",
                passed=True,
                message="expected_titles下的正文字体格式检查通过",
                details={"location": "expected_titles下的正文段落"}
            ))
            
        return results
        
    def _is_paragraph_under_expected_title(self, para_idx: int, expected_titles: List[str]) -> bool:
        """
        检查段落是否在expected_titles下面
        
        Args:
            para_idx: 段落索引（从1开始）
            expected_titles: 期望的标题文本列表
        
        Returns:
            bool: 是否在expected_titles下面
        """
        # 查找当前段落之前的最近标题
        current_title_text = None
        for i in range(para_idx - 1, 0, -1):
            para = self.doc.paragraphs[i - 1]  # 索引从0开始，但para_idx从1开始
            style_name = get_paragraph_style_name(para)
            
            # 如果是标题段落
            if style_name and style_name.startswith(("Heading", "标题")):
                current_title_text = para.text.strip()
                break
        
        # 如果找不到标题，则不在expected_titles下面
        if not current_title_text:
            return False
            
        # 检查标题是否在expected_titles中
        # 使用模糊匹配，允许标题中包含额外的字符（如编号）
        for expected_title in expected_titles:
            # 如果期望的标题是当前标题的一部分，或者当前标题是期望标题的一部分
            if expected_title in current_title_text or current_title_text in expected_title:
                return True
                
        return False
        
    def _find_parent_title(self, para_idx: int) -> Optional[str]:
        """
        查找段落所属的标题
        
        Args:
            para_idx: 段落索引（从1开始）
        
        Returns:
            str: 标题文本，如果找不到则返回None
        """
        # 查找当前段落之前的最近标题
        for i in range(para_idx - 1, 0, -1):
            para = self.doc.paragraphs[i - 1]  # 索引从0开始，但para_idx从1开始
            style_name = get_paragraph_style_name(para)
            
            # 如果是标题段落
            if style_name and style_name.startswith(("Heading", "标题")):
                return para.text.strip()
                
        return None
        
    def check_fonts(self) -> List[CheckResult]:
        """检查所有字体"""
        # 检查标题字体
        heading_results = self.check_heading_fonts()
        
        # 检查正文字体
        content_results = self.check_content_fonts()
        
        # 合并结果
        results = []
        results.extend(heading_results)
        results.extend(content_results)
        
        return results