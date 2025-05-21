"""检查器测试模块"""
from typing import List
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import logging

def debug_print_document(doc_path: str) -> None:
    """调试打印文档的所有内容，包括标题、段落和表格"""
    doc = Document(doc_path)
    
    print("\n=== 文档结构分析 ===\n")
    
    print("【1. 所有段落及其样式】")
    print("-" * 50)
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:  # 只打印非空段落
            style = para.style.name if para.style else "无样式"
            print(f"段落 {i}:")
            print(f"  样式: {style}")
            print(f"  内容: {text}")
            print(f"  长度: {len(text)} 字符")
            print("-" * 30)
    
    print("\n【2. 所有标题】")
    print("-" * 50)
    for i, para in enumerate(doc.paragraphs, 1):
        style = para.style.name if para.style else "无样式"
        if style and style.startswith(("Heading", "标题")):
            print(f"标题 {i}:")
            print(f"  级别: {style}")
            print(f"  内容: {para.text.strip()}")
            print("-" * 30)
    
    print("\n【3. 所有表格】")
    print("-" * 50)
    for i, table in enumerate(doc.tables, 1):
        print(f"表格 {i}:")
        print(f"  行数: {len(table.rows)}")
        print(f"  列数: {len(table.columns)}")
        print("  内容:")
        for row_idx, row in enumerate(table.rows, 1):
            row_content = " | ".join(cell.text.strip() for cell in row.cells)
            print(f"    行{row_idx}: {row_content}")
        print("-" * 30)

def find_paragraphs_after_heading(doc: Document, heading_text: str) -> List[Paragraph]:
    """调试用：查找指定标题后的段落"""
    results = []
    heading_found = False
    
    for para in doc.paragraphs:
        if not heading_found:
            if para.text.strip() == heading_text:
                heading_found = True
                print(f"\n找到标题: '{heading_text}'")
            continue
        
        if para.style.name and para.style.name.startswith(("Heading", "标题")):
            break
        
        text = para.text.strip()
        if text:
            results.append(para)
            print(f"  后续段落: {text}")
    
    if not heading_found:
        print(f"\n未找到标题: '{heading_text}'")
    elif not results:
        print("  未找到后续段落")
    
    return results

def find_tables_under_heading(doc: Document, heading_text: str) -> List[Table]:
    """调试用：查找指定标题下的表格"""
    tables = []
    heading_found = False
    next_heading_found = False
    
    for i, para in enumerate(doc.paragraphs):
        # 查找标题
        if not heading_found:
            if para.text.strip() == heading_text:
                heading_found = True
                print(f"\n找到标题: '{heading_text}'")
            continue
        
        # 检查是否到达下一个标题
        if heading_found and not next_heading_found:
            if para.style.name and para.style.name.startswith(("Heading", "标题")):
                next_heading_found = True
                continue
        
        # 在两个标题之间查找表格
        try:
            element = doc._element.body.getchildren()[i]
            if element.tag.endswith('tbl'):
                table = Table(element, doc)
                tables.append(table)
                print(f"  找到表格: {len(table.rows)}行 x {len(table.columns)}列")
                # 打印表格内容预览
                for row_idx, row in enumerate(table.rows[:3], 1):  # 只显示前3行
                    row_content = " | ".join(cell.text.strip() for cell in row.cells)
                    print(f"    行{row_idx}: {row_content}")
                if len(table.rows) > 3:
                    print("    ...")
        except Exception as e:
            continue
    
    if not heading_found:
        print(f"\n未找到标题: '{heading_text}'")
    elif not tables:
        print("  未找到表格")
    
    return tables

def main():
    """测试主函数"""
    import sys
    if len(sys.argv) < 2:
        print("使用方法: python test_checker.py <word文档路径>")
        sys.exit(1)
    
    doc_path = sys.argv[1]
    try:
        print(f"\n开始分析文档: {doc_path}")
        print("=" * 60)
        
        # # 1. 打印整个文档结构
        # debug_print_document(doc_path)
        
        # 2. 测试标题后的段落查找
        doc = Document(doc_path)
        print("\n\n=== 测试标题后段落查找 ===")
        for heading in ["测试工具版本", "适用产品", "BIOS设置", "OS设置", "环境部署", "测试执行"]:
            find_paragraphs_after_heading(doc, heading)
        
        # 3. 测试标题下的表格查找
        print("\n\n=== 测试标题下表格查找 ===")
        for heading in ["测试工具版本", "适用产品"]:
            find_tables_under_heading(doc, heading)
            
    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
